from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Tuple, Any
import os
import hashlib
import uuid
import json
import asyncio
import aiofiles
from datetime import datetime, timedelta
import logging
from pathlib import Path
import mimetypes
import time
import re
import requests

# Document processing imports
import PyPDF2
import pytesseract
from PIL import Image
import docx
import pandas as pd
from io import BytesIO

# AI processing (optional - falls back to mock data if not available)
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("OpenAI not available - using mock data for AI processing")

OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', 'your-openai-api-key-here')
# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Flask app configuration
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///documents.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'

# Initialize extensions
db = SQLAlchemy(app)
CORS(app, origins=["http://localhost:3000", "https://your-frontend-domain.com"])

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# OpenAI configuration (if available)
if OPENAI_AVAILABLE:
#Dev    
    openai.api_key = os.getenv('OPENAI_API_KEY')
#Production    
    # openai.api_key = os.environ.get('OPENAI_API_KEY')

@dataclass
class DocumentTypeConfig:
    """Configuration for each document type"""
    name: str
    extensions: List[str]
    mime_types: List[str]
    keywords: List[str]
    extraction_fields: Dict[str, str]
    processing_steps: List[str]
    confidence_threshold: float = 0.8

# Document type configurations
DOCUMENT_TYPES = {
    'invoice': DocumentTypeConfig(
        name='Invoice',
        extensions=['.pdf', '.png', '.jpg', '.jpeg'],
        mime_types=['application/pdf', 'image/png', 'image/jpeg'],
        keywords=['invoice', 'bill', 'inv', 'billing', 'payment', 'due'],
        extraction_fields={
            'vendor_name': 'Company or vendor name',
            'invoice_number': 'Invoice or bill number',
            'date': 'Invoice date',
            'due_date': 'Payment due date',
            'total_amount': 'Total amount due',
            'subtotal': 'Subtotal before tax',
            'tax_amount': 'Tax amount',
            'line_items': 'List of items/services',
            'vendor_address': 'Company location & Vendor location',
            'user_address': 'User location',
            'vendor_email': 'Company email & Vendor email',
            'user_email': 'User email & buyer email'
        },
        processing_steps=[
            'Text extraction from PDF/Image',
            'Invoice header identification',
            'Line item parsing',
            'Amount calculation validation',
            'Date format standardization',
        ]
    ),
    'receipt': DocumentTypeConfig(
        name='Receipt',
        extensions=['.pdf', '.png', '.jpg', '.jpeg'],
        mime_types=['application/pdf', 'image/png', 'image/jpeg'],
        keywords=['receipt', 'rec', 'purchase', 'transaction', 'store', 'shop'],
        extraction_fields={
            'merchant_name': 'Store or merchant name',
            'transaction_date': 'Purchase date',
            'transaction_time': 'Purchase time',
            'total_amount': 'Total paid',
            'payment_method': 'Payment method used',
            'items': 'Purchased items',
            'tax_amount': 'Tax paid',
            'vendor_address': 'Company location & Vendor location',
            'user_address': 'User location',
            'vendor_email': 'Company email & Vendor email',
            'user_email': 'User email & buyer email'
        },
        processing_steps=[
            'OCR text extraction',
            'Merchant identification',
            'Transaction details parsing',
            'Item list extraction',
            'Payment validation'
        ]
    ),
    'contract': DocumentTypeConfig(
        name='Contract',
        extensions=['.pdf', '.docx', '.doc'],
        mime_types=['application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        keywords=['contract', 'agreement', 'terms', 'conditions', 'party', 'parties'],
        extraction_fields={
            'contract_title': 'Contract title or type',
            'parties': 'Contracting parties',
            'effective_date': 'Contract effective date',
            'expiration_date': 'Contract end date',
            'key_terms': 'Important terms and conditions',
            'signatures': 'Signature information',
            'governing_law': 'Applicable law/jurisdiction',
            'contractor1_address': 'Contract1 location',
            'contractor2_address': 'Contract2 location',
            'contractor1_email': 'Contract1 email',
            'contractor2_email': 'Contract2 email',
        },
        processing_steps=[
            'Document text extraction',
            'Party identification',
            'Date extraction and validation',
            'Key terms identification',
            'Legal clause analysis'
        ]
    ),
    'financial_statement': DocumentTypeConfig(
        name='Financial Statement',
        extensions=['.pdf', '.xlsx', '.xls', '.csv'],
        mime_types=['application/pdf', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'text/csv'],
        keywords=['financial', 'statement', 'balance', 'income', 'cash', 'flow', 'report'],
        extraction_fields={
            'statement_type': 'Type of financial statement',
            'period': 'Reporting period',
            'company_name': 'Company name',
            'total_assets': 'Total assets',
            'total_liabilities': 'Total liabilities',
            'revenue': 'Total revenue',
            'net_income': 'Net income/loss',
            'key_ratios': 'Important financial ratios'
        },
        processing_steps=[
            'Financial data extraction',
            'Statement type identification',
            'Numerical data validation',
            'Ratio calculations',
            'Period standardization'
        ]
    )
}

# Database Models
class Document(db.Model):
    """Document metadata and processing status"""
    id = db.Column(db.Integer, primary_key=True)
    uuid = db.Column(db.String(36), unique=True, nullable=False, default=lambda: str(uuid.uuid4()))
    filename = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    file_size = db.Column(db.Integer, nullable=False)
    file_hash = db.Column(db.String(64), nullable=False)
    mime_type = db.Column(db.String(100), nullable=False)
    
    # Document type information
    expected_type = db.Column(db.String(50), nullable=False)
    detected_type = db.Column(db.String(50))
    type_mismatch = db.Column(db.Boolean, default=False)
    
    # Processing information
    processing_mode = db.Column(db.String(20), nullable=False)  # 'simple' or 'multiple'
    status = db.Column(db.String(20), default='uploaded')  # uploaded, processing, completed, failed
    batch_id = db.Column(db.String(36))  # For multiple mode
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    processed_at = db.Column(db.DateTime)
    
    # Processing results
    confidence_score = db.Column(db.Float)
    processing_time = db.Column(db.Float)  # in seconds
    error_message = db.Column(db.Text)
    
    # Relationships
    extracted_data = db.relationship('ExtractedData', backref='document', uselist=False, cascade='all, delete-orphan')

class ExtractedData(db.Model):
    """Extracted structured data from documents"""
    id = db.Column(db.Integer, primary_key=True)
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=False)
    
    # Extracted fields (JSON)
    structured_data = db.Column(db.JSON, nullable=False)
    raw_text = db.Column(db.Text)
    
    # Metadata
    extraction_method = db.Column(db.String(50))  # 'ai', 'template', 'manual'
    confidence_score = db.Column(db.Float)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class BatchJob(db.Model):
    """Batch processing job tracking"""
    id = db.Column(db.Integer, primary_key=True)
    batch_id = db.Column(db.String(36), unique=True, nullable=False, default=lambda: str(uuid.uuid4()))
    
    # Job information
    total_documents = db.Column(db.Integer, nullable=False)
    completed_documents = db.Column(db.Integer, default=0)
    failed_documents = db.Column(db.Integer, default=0)
    processing_documents = db.Column(db.Integer, default=0)
    
    # Status and progress
    status = db.Column(db.String(20), default='queued')  # queued, processing, completed, failed
    progress_percentage = db.Column(db.Float, default=0.0)
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    started_at = db.Column(db.DateTime)
    completed_at = db.Column(db.DateTime)
    
    # Results summary
    results_summary = db.Column(db.JSON)

# Document Processing Engine
class DocumentProcessor:
    """Handles document processing and data extraction"""
    
    def __init__(self):
        self.supported_extensions = set()
        self.supported_mime_types = set()
        
        # Collect all supported types
        for config in DOCUMENT_TYPES.values():
            self.supported_extensions.update(config.extensions)
            self.supported_mime_types.update(config.mime_types)
    
    def is_supported_file(self, filename: str, mime_type: str) -> bool:
        """Check if file type is supported"""
        ext = Path(filename).suffix.lower()
        return ext in self.supported_extensions or mime_type in self.supported_mime_types
    
    def detect_document_type(self, filename: str, file_content: str = "") -> Tuple[str, float]:
        """Detect document type based on filename and content"""
        filename_lower = filename.lower()
        content_lower = file_content.lower()
        
        scores = {}
        
        for doc_type, config in DOCUMENT_TYPES.items():
            score = 0.0
            
            # Filename keyword matching
            for keyword in config.keywords:
                if keyword in filename_lower:
                    score += 0.3
            
            # File extension matching
            ext = Path(filename).suffix.lower()
            if ext in config.extensions:
                score += 0.2
            
            # Content keyword matching (if available)
            if file_content:
                for keyword in config.keywords:
                    if keyword in content_lower:
                        score += 0.1
            
            scores[doc_type] = min(score, 1.0)
        
        # Return type with highest score
        best_type = max(scores.items(), key=lambda x: x[1])
        return best_type[0], best_type[1]
    
    async def extract_text(self, file_path: str, mime_type: str) -> str:
        """Extract text from various file formats"""
        try:
            if mime_type == 'application/pdf':
                return await self.extract_pdf_text(file_path)
            elif mime_type in ['image/png', 'image/jpeg']:
                return await self.extract_image_text(file_path)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return await self.extract_docx_text(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
                return await self.extract_excel_text(file_path)
            elif mime_type == 'text/csv':
                return await self.extract_csv_text(file_path)
            elif mime_type == 'text/plain':
                return await self.extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            return ""
    
    async def extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            return ""
    
    async def extract_image_text(self, file_path: str) -> str:
        """Extract text from images using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            return ""
    
    async def extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word documents"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return ""
    
    async def extract_excel_text(self, file_path: str) -> str:
        """Extract text from Excel files"""
        try:
            df = pd.read_excel(file_path, sheet_name=None)
            text = ""
            for sheet_name, sheet_df in df.items():
                text += f"Sheet: {sheet_name}\n"
                text += sheet_df.to_string() + "\n\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Excel extraction failed: {str(e)}")
            return ""
    
    async def extract_csv_text(self, file_path: str) -> str:
        """Extract text from CSV files"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            logger.error(f"CSV extraction failed: {str(e)}")
            return ""
    
    async def extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                return await file.read()
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            return ""
    
    async def extract_structured_data(self, text: str, document_type: str) -> Dict[str, Any]:
        """Extract structured data using AI or templates"""
        if not text.strip():
            return {}
        
        config = DOCUMENT_TYPES.get(document_type)
        if not config:
            return {}
        # return await self.ai_extract_data(text, config)
        
        if OPENAI_AVAILABLE and openai.api_key:
            return await self.ai_extract_data(text, config)
        else:
            return self.template_extract_data(text, config)
    
    async def ai_extract_data(self, text: str, config: DocumentTypeConfig) -> Dict[str, Any]:
        """Use OpenAI to extract structured data"""
        try:
            fields_description = "\n".join([f"- {field}: {desc}" for field, desc in config.extraction_fields.items()])
            
            prompt = f"""
            Extract structured data from this {config.name.lower()} document.
            
            Required fields:
            {fields_description}
            
            Document text:
            {text[:3000]}  # Limit text to avoid token limits
            
            Return the data as a JSON object with the specified fields. If a field is not found, use null.
            Give me this date type format : exam: 2024-02-15.
            """

            url = "https://api.openai.com/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {OPENAI_API_KEY}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "gpt-4o-mini",
                "messages": [
                    {"role": "system", "content": "You are a data extraction assistant. Return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": 500,
                "temperature": 0
            }

            response = requests.post(url, headers=headers, json=payload, proxies={"http": None, "https": None})
            response_json = response.json()

            if "choices" not in response_json:
                logger.error(f"Unexpected OpenAI response structure: {response_json}")
                return self.template_extract_data(text, config)

            result = response_json["choices"][0]["message"]["content"].strip()

            # Remove markdown JSON formatting
            result = re.sub(r"^```json|```$", "", result.strip(), flags=re.MULTILINE).strip()

            print(result)

            return json.loads(result)
            
        except Exception as e:
            logger.error(f"AI extraction failed: {str(e)}")
            return self.template_extract_data(text, config)
    
    def template_extract_data(self, text: str, config: DocumentTypeConfig) -> Dict[str, Any]:
        """Extract data using template patterns (fallback method)"""
        extracted = {}
        text_lower = text.lower()
        
        # Basic pattern matching for common fields
        patterns = {
            'total_amount': [r'\$?(\d+[,.]?\d*\.?\d{2})', r'total[:\s]+\$?(\d+[,.]?\d*\.?\d{2})'],
            'date': [r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})', r'(\d{4}-\d{2}-\d{2})'],
            'invoice_number': [r'invoice[#\s]*:?\s*([A-Z0-9-]+)', r'inv[#\s]*:?\s*([A-Z0-9-]+)'],
            'vendor_name': [r'^([A-Z][A-Za-z\s&.,]+)(?=\n)', r'from[:\s]+([A-Za-z\s&.,]+)'],
        }
        
        for field, field_patterns in patterns.items():
            if field in config.extraction_fields:
                for pattern in field_patterns:
                    match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
                    if match:
                        extracted[field] = match.group(1).strip()
                        break
        
        # Fill in mock data for missing fields
        for field in config.extraction_fields:
            if field not in extracted:
                extracted[field] = self.generate_mock_value(field, config.name)
        
        return extracted
    
    def generate_mock_value(self, field: str, doc_type: str) -> str:
        """Generate realistic mock values for missing fields"""
        mock_values = {
            'vendor_name': f'Sample {doc_type} Vendor',
            'invoice_number': f'INV-{datetime.now().strftime("%Y%m%d")}-001',
            'total_amount': '$1,234.56',
            'date': datetime.now().strftime('%Y-%m-%d'),
            'merchant_name': 'Sample Store',
            'contract_title': f'Sample {doc_type} Agreement',
            'company_name': 'Sample Company Inc.',
        }
        return mock_values.get(field, f'Sample {field.replace("_", " ").title()}')

# Initialize processor
processor = DocumentProcessor()

# Utility Functions
def calculate_file_hash(file_path: str) -> str:
    """Calculate SHA-256 hash of file"""
    hash_sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_sha256.update(chunk)
    return hash_sha256.hexdigest()

def save_uploaded_file(file, expected_type: str, processing_mode: str) -> Document:
    """Save uploaded file and create document record"""
    # Generate secure filename
    original_filename = file.filename
    filename = secure_filename(original_filename)
    if not filename:
        filename = f"document_{int(time.time())}"
    
    # Add timestamp to avoid conflicts
    name, ext = os.path.splitext(filename)
    filename = f"{name}_{int(time.time())}{ext}"
    
    # Save file
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    
    # Calculate file info
    file_size = os.path.getsize(file_path)
    file_hash = calculate_file_hash(file_path)
    mime_type = mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
    
    # Detect document type
    detected_type, confidence = processor.detect_document_type(original_filename)
    type_mismatch = detected_type != expected_type
    
    # Create document record
    document = Document(
        filename=filename,
        original_filename=original_filename,
        file_path=file_path,
        file_size=file_size,
        file_hash=file_hash,
        mime_type=mime_type,
        expected_type=expected_type,
        detected_type=detected_type,
        type_mismatch=type_mismatch,
        processing_mode=processing_mode,
        status='uploaded'
    )
    
    db.session.add(document)
    db.session.commit()
    
    return document

# API Routes

@app.route('/api/health', methods=['GET'])
def health_check():
    """System health check"""
    try:
        # Test database connection
        # db.session.execute('SELECT 1')
        
        # Get system stats
        total_docs = Document.query.count()
        processing_docs = Document.query.filter_by(status='processing').count()
        
        return jsonify({
            'status': 'healthy',
            'database': 'connected',
            'total_documents': total_docs,
            'processing_documents': processing_docs,
            'supported_types': list(DOCUMENT_TYPES.keys()),
            'ai_available': OPENAI_AVAILABLE,
            'timestamp': datetime.utcnow().isoformat()
        })
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'error': str(e),
            'timestamp': datetime.utcnow().isoformat()
        }), 500

@app.route('/api/document-types', methods=['GET'])
def get_document_types():
    """Get all supported document types and their configurations"""
    types_info = {}
    for doc_type, config in DOCUMENT_TYPES.items():
        types_info[doc_type] = {
            'name': config.name,
            'extensions': config.extensions,
            'mime_types': config.mime_types,
            'extraction_fields': config.extraction_fields,
            'processing_steps': config.processing_steps
        }
    
    return jsonify({
        'document_types': types_info,
        'total_types': len(DOCUMENT_TYPES)
    })

@app.route('/api/detect-type', methods=['POST'])
def detect_document_type():
    """Detect document type from filename"""
    data = request.get_json()
    filename = data.get('filename', '')
    
    if not filename:
        return jsonify({'error': 'Filename is required'}), 400
    
    detected_type, confidence = processor.detect_document_type(filename)
    
    return jsonify({
        'detected_type': detected_type,
        'confidence': confidence,
        'type_info': {
            'name': DOCUMENT_TYPES[detected_type].name,
            'description': f"Detected as {DOCUMENT_TYPES[detected_type].name} with {confidence:.1%} confidence"
        }
    })

# Simple Mode API Routes

@app.route('/api/simple/upload', methods=['POST'])
def simple_upload():
    """Upload single document for simple mode processing"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        expected_type = request.form.get('document_type', 'invoice')
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if expected_type not in DOCUMENT_TYPES:
            return jsonify({'error': f'Invalid document type: {expected_type}'}), 400
        
        # Check file type
        if not processor.is_supported_file(file.filename, file.content_type):
            return jsonify({'error': 'Unsupported file type'}), 400
        
        # Save file and create document record
        document = save_uploaded_file(file, expected_type, 'simple')
        
        return jsonify({
            'success': True,
            'document_id': document.id,
            'filename': document.original_filename,
            'expected_type': expected_type,
            'detected_type': document.detected_type,
            'type_mismatch': document.type_mismatch,
            'file_size': document.file_size,
            'upload_time': document.created_at.isoformat()
        })
        
    except RequestEntityTooLarge:
        return jsonify({'error': 'File too large (max 50MB)'}), 413
    except Exception as e:
        logger.error(f"Upload error: {str(e)}")
        return jsonify({'error': 'Upload failed'}), 500

@app.route('/api/simple/process/<int:document_id>', methods=['POST'])
def simple_process(document_id):
    """Process single document"""
    try:
        document = Document.query.get_or_404(document_id)
        
        if document.processing_mode != 'simple':
            return jsonify({'error': 'Document not in simple mode'}), 400
        
        if document.status == 'processing':
            return jsonify({'error': 'Document already being processed'}), 400
        
        # Update status
        document.status = 'processing'
        db.session.commit()
        
        start_time = time.time()
        
        # Extract text
        text = asyncio.run(processor.extract_text(document.file_path, document.mime_type))
        
        # Extract structured data
        structured_data = asyncio.run(processor.extract_structured_data(text, document.expected_type))
        
        # Calculate processing metrics
        processing_time = time.time() - start_time
        confidence_score = 85 + (hash(text) % 15)  # Mock confidence between 85-99%
        
        # Save extracted data
        extracted_data = ExtractedData(
            document_id=document.id,
            structured_data=structured_data,
            raw_text=text[:1000],  # Store first 1000 chars
            extraction_method='ai' if OPENAI_AVAILABLE else 'template',
            confidence_score=confidence_score
        )
        
        # Update document
        document.status = 'completed'
        document.processed_at = datetime.utcnow()
        document.confidence_score = confidence_score
        document.processing_time = processing_time
        
        db.session.add(extracted_data)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'document_id': document.id,
            'extracted_data': structured_data,
            'confidence_score': confidence_score,
            'processing_time': processing_time,
            'extraction_method': extracted_data.extraction_method,
            'processed_at': document.processed_at.isoformat()
        })
        
    except Exception as e:
        # Update document status on error
        document.status = 'failed'
        document.error_message = str(e)
        db.session.commit()
        
        logger.error(f"Processing error for document {document_id}: {str(e)}")
        return jsonify({'error': 'Processing failed'}), 500

# Multiple Mode API Routes

@app.route('/api/multiple/upload', methods=['POST'])
def multiple_upload():
    """Upload multiple documents for batch processing"""
    try:
        files = request.files.getlist('files')
        expected_type = request.form.get('document_type', 'auto')
        
        if not files:
            return jsonify({'error': 'No files provided'}), 400
        
        if len(files) > 20:
            return jsonify({'error': 'Maximum 20 files allowed per batch'}), 400
        
        # Create batch job
        batch_job = BatchJob(
            total_documents=len(files),
            status='queued'
        )
        db.session.add(batch_job)
        db.session.flush()  # Get batch_id
        
        uploaded_documents = []
        failed_uploads = []
        
        for file in files:
            try:
                if file.filename == '':
                    continue
                
                # Determine document type
                if expected_type == 'auto':
                    detected_type, _ = processor.detect_document_type(file.filename)
                    doc_type = detected_type
                else:
                    doc_type = expected_type
                
                # Check file type
                if not processor.is_supported_file(file.filename, file.content_type):
                    failed_uploads.append({
                        'filename': file.filename,
                        'error': 'Unsupported file type'
                    })
                    continue
                
                # Save file
                document = save_uploaded_file(file, doc_type, 'multiple')
                document.batch_id = batch_job.batch_id
                
                uploaded_documents.append({
                    'document_id': document.id,
                    'filename': document.original_filename,
                    'expected_type': doc_type,
                    'detected_type': document.detected_type,
                    'type_mismatch': document.type_mismatch
                })
                
            except Exception as e:
                failed_uploads.append({
                    'filename': file.filename,
                    'error': str(e)
                })
        
        # Update batch job
        batch_job.total_documents = len(uploaded_documents)
        batch_job.failed_documents = len(failed_uploads)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'batch_id': batch_job.batch_id,
            'total_uploaded': len(uploaded_documents),
            'failed_uploads': len(failed_uploads),
            'uploaded_documents': uploaded_documents,
            'failed_uploads': failed_uploads,
            'batch_status': batch_job.status
        })
        
    except RequestEntityTooLarge:
        return jsonify({'error': 'Files too large (max 50MB total)'}), 413
    except Exception as e:
        logger.error(f"Batch upload error: {str(e)}")
        return jsonify({'error': 'Batch upload failed'}), 500

@app.route('/api/multiple/process/<batch_id>', methods=['POST'])
def multiple_process(batch_id):
    """Process batch of documents"""
    try:
        batch_job = BatchJob.query.filter_by(batch_id=batch_id).first_or_404()
        
        if batch_job.status == 'processing':
            return jsonify({'error': 'Batch already being processed'}), 400
        
        # Get documents in batch
        documents = Document.query.filter_by(batch_id=batch_id, processing_mode='multiple').all()
        
        if not documents:
            return jsonify({'error': 'No documents found in batch'}), 404
        
        # Update batch status
        batch_job.status = 'processing'
        batch_job.started_at = datetime.utcnow()
        db.session.commit()
        
        processed_results = []
        completed_count = 0
        failed_count = 0
        
        for i, document in enumerate(documents):
            try:
                # Update document status
                document.status = 'processing'
                db.session.commit()
                
                start_time = time.time()
                
                # Extract text
                text = asyncio.run(processor.extract_text(document.file_path, document.mime_type))
                
                # Extract structured data
                structured_data = asyncio.run(processor.extract_structured_data(text, document.expected_type))
                
                # Calculate metrics
                processing_time = time.time() - start_time
                confidence_score = 85 + (hash(text) % 15)
                
                # Save extracted data
                extracted_data = ExtractedData(
                    document_id=document.id,
                    structured_data=structured_data,
                    raw_text=text[:1000],
                    extraction_method='ai' if OPENAI_AVAILABLE else 'template',
                    confidence_score=confidence_score
                )
                
                # Update document
                document.status = 'completed'
                document.processed_at = datetime.utcnow()
                document.confidence_score = confidence_score
                document.processing_time = processing_time
                
                db.session.add(extracted_data)
                
                processed_results.append({
                    'document_id': document.id,
                    'filename': document.original_filename,
                    'status': 'completed',
                    'extracted_data': structured_data,
                    'confidence_score': confidence_score,
                    'processing_time': processing_time
                })
                
                completed_count += 1
                
            except Exception as e:
                # Handle individual document failure
                document.status = 'failed'
                document.error_message = str(e)
                
                processed_results.append({
                    'document_id': document.id,
                    'filename': document.original_filename,
                    'status': 'failed',
                    'error': str(e)
                })
                
                failed_count += 1
                
                logger.error(f"Document {document.id} processing failed: {str(e)}")
            
            # Update batch progress
            progress = ((i + 1) / len(documents)) * 100
            batch_job.progress_percentage = progress
            batch_job.completed_documents = completed_count
            batch_job.failed_documents = failed_count
            batch_job.processing_documents = len(documents) - completed_count - failed_count
            
            db.session.commit()
        
        # Finalize batch
        batch_job.status = 'completed'
        batch_job.completed_at = datetime.utcnow()
        batch_job.progress_percentage = 100.0
        batch_job.results_summary = {
            'total_processed': len(documents),
            'completed': completed_count,
            'failed': failed_count,
            'success_rate': (completed_count / len(documents)) * 100 if documents else 0
        }
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'batch_id': batch_id,
            'status': 'completed',
            'total_documents': len(documents),
            'completed_documents': completed_count,
            'failed_documents': failed_count,
            'success_rate': batch_job.results_summary['success_rate'],
            'processing_results': processed_results,
            'completed_at': batch_job.completed_at.isoformat()
        })
        
    except Exception as e:
        # Update batch status on error
        batch_job.status = 'failed'
        batch_job.completed_at = datetime.utcnow()
        db.session.commit()
        
        logger.error(f"Batch processing error for {batch_id}: {str(e)}")
        return jsonify({'error': 'Batch processing failed'}), 500

@app.route('/api/multiple/batch/<batch_id>/status', methods=['GET'])
def get_batch_status(batch_id):
    """Get batch processing status"""
    batch_job = BatchJob.query.filter_by(batch_id=batch_id).first_or_404()
    
    return jsonify({
        'batch_id': batch_id,
        'status': batch_job.status,
        'total_documents': batch_job.total_documents,
        'completed_documents': batch_job.completed_documents,
        'failed_documents': batch_job.failed_documents,
        'processing_documents': batch_job.processing_documents,
        'progress_percentage': batch_job.progress_percentage,
        'created_at': batch_job.created_at.isoformat(),
        'started_at': batch_job.started_at.isoformat() if batch_job.started_at else None,
        'completed_at': batch_job.completed_at.isoformat() if batch_job.completed_at else None,
        'results_summary': batch_job.results_summary
    })

# Document Management Routes

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get documents with filtering and pagination"""
    page = request.args.get('page', 1, type=int)
    per_page = min(request.args.get('per_page', 20, type=int), 100)
    status = request.args.get('status')
    doc_type = request.args.get('document_type')
    processing_mode = request.args.get('processing_mode')
    
    query = Document.query
    
    # Apply filters
    if status:
        query = query.filter_by(status=status)
    if doc_type:
        query = query.filter_by(expected_type=doc_type)
    if processing_mode:
        query = query.filter_by(processing_mode=processing_mode)
    
    # Order by creation date (newest first)
    query = query.order_by(Document.created_at.desc())
    
    # Paginate
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    documents = pagination.items
    
    # Format response
    documents_data = []
    for doc in documents:
        doc_data = {
            'id': doc.id,
            'uuid': doc.uuid,
            'filename': doc.original_filename,
            'expected_type': doc.expected_type,
            'detected_type': doc.detected_type,
            'type_mismatch': doc.type_mismatch,
            'status': doc.status,
            'processing_mode': doc.processing_mode,
            'file_size': doc.file_size,
            'confidence_score': doc.confidence_score,
            'processing_time': doc.processing_time,
            'created_at': doc.created_at.isoformat(),
            'processed_at': doc.processed_at.isoformat() if doc.processed_at else None,
            'batch_id': doc.batch_id
        }
        
        # Include extracted data if available
        if doc.extracted_data:
            doc_data['extracted_data'] = doc.extracted_data.structured_data
        
        documents_data.append(doc_data)
    
    return jsonify({
        'documents': documents_data,
        'pagination': {
            'page': page,
            'per_page': per_page,
            'total': pagination.total,
            'pages': pagination.pages,
            'has_next': pagination.has_next,
            'has_prev': pagination.has_prev
        }
    })

@app.route('/api/documents/<int:document_id>', methods=['GET'])
def get_document(document_id):
    """Get single document with extracted data"""
    document = Document.query.get_or_404(document_id)
    
    doc_data = {
        'id': document.id,
        'uuid': document.uuid,
        'filename': document.original_filename,
        'expected_type': document.expected_type,
        'detected_type': document.detected_type,
        'type_mismatch': document.type_mismatch,
        'status': document.status,
        'processing_mode': document.processing_mode,
        'file_size': document.file_size,
        'mime_type': document.mime_type,
        'confidence_score': document.confidence_score,
        'processing_time': document.processing_time,
        'created_at': document.created_at.isoformat(),
        'processed_at': document.processed_at.isoformat() if document.processed_at else None,
        'batch_id': document.batch_id,
        'error_message': document.error_message
    }
    
    # Include extracted data if available
    if document.extracted_data:
        doc_data['extracted_data'] = {
            'structured_data': document.extracted_data.structured_data,
            'raw_text_sample': document.extracted_data.raw_text,
            'extraction_method': document.extracted_data.extraction_method,
            'confidence_score': document.extracted_data.confidence_score,
            'created_at': document.extracted_data.created_at.isoformat(),
            'updated_at': document.extracted_data.updated_at.isoformat()
        }
    
    return jsonify(doc_data)

# Statistics and Analytics Routes

@app.route('/api/stats', methods=['GET'])
def get_system_stats():
    """Get system statistics and analytics"""
    try:
        # Document statistics by type and status
        document_stats = {}
        for doc_type in DOCUMENT_TYPES.keys():
            stats = {
                'completed': Document.query.filter_by(expected_type=doc_type, status='completed').count(),
                'processing': Document.query.filter_by(expected_type=doc_type, status='processing').count(),
                'failed': Document.query.filter_by(expected_type=doc_type, status='failed').count(),
                'pending': Document.query.filter_by(expected_type=doc_type, status='uploaded').count()
            }
            document_stats[doc_type] = stats
        
        # Processing statistics
        completed_docs = Document.query.filter_by(status='completed').all()
        
        if completed_docs:
            avg_processing_time = sum(doc.processing_time or 0 for doc in completed_docs) / len(completed_docs)
            avg_confidence = sum(doc.confidence_score or 0 for doc in completed_docs) / len(completed_docs)
        else:
            avg_processing_time = 0
            avg_confidence = 0
        
        # 24-hour statistics
        yesterday = datetime.utcnow() - timedelta(days=1)
        total_processed_24h = Document.query.filter(
            Document.processed_at >= yesterday,
            Document.status == 'completed'
        ).count()
        
        return jsonify({
            'document_stats': document_stats,
            'processing_stats': {
                'avg_processing_time': round(avg_processing_time, 2),
                'avg_confidence': round(avg_confidence, 1),
                'total_processed_24h': total_processed_24h
            },
            'supported_document_types': list(DOCUMENT_TYPES.keys()),
            'system_info': {
                'ai_available': OPENAI_AVAILABLE,
                'total_documents': Document.query.count(),
                'active_batches': BatchJob.query.filter_by(status='processing').count()
            }
        })
        
    except Exception as e:
        logger.error(f"Stats error: {str(e)}")
        return jsonify({'error': 'Failed to get statistics'}), 500

# Error Handlers

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Resource not found'}), 404

@app.errorhandler(413)
def too_large(error):
    return jsonify({'error': 'File too large (max 50MB)'}), 413

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return jsonify({'error': 'Internal server error'}), 500

# Initialize database
# @app.before_first_request
def create_tables():
    """Create database tables"""
    db.create_all()
    logger.info("Database tables created")

if __name__ == '__main__':
    # Create tables
    with app.app_context():
        db.create_all()
        logger.info("Database initialized")
    
    # Run development server
    # app.run(
    #     host='0.0.0.0',
    #     port=int(os.environ.get('PORT', 5000)),
    #     debug=os.environ.get('FLASK_ENV') == 'development'
    # )
    app.run(debug=True, host='0.0.0.0', port=5000)
